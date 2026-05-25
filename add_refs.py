#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Add missing references to Folliscope_IEEE.docx"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

# Load dokumen
doc = Document('Folliscope_IEEE.docx')

# Find REFERENSI section
ref_para_index = None
for i, para in enumerate(doc.paragraphs):
    if 'REFERENSI' in para.text.upper():
        ref_para_index = i
        print(f"[OK] Found REFERENSI at paragraph index {i}")
        break

if ref_para_index is None:
    print("[ERROR] REFERENSI section not found!")
    exit(1)

# Missing references (yang ada di README tapi tidak di DOCX)
missing_refs = [
    ("3", "Choong CS, et al. (1996). Reduced androgen receptor gene expression with first exon CAG repeat expansion. Mol Endocrinol, 10(12):1527-1535."),
    ("4", "Ellis JA, et al. (2001). Polymorphism of the androgen receptor gene is associated with male pattern baldness. J Invest Dermatol, 116(3):452-455."),
    ("6", "Norwood OT. (1975). Male pattern baldness: Classification and incidence. South Med J, 68(11):1359-1365."),
    ("7", "Ludwig E. (1977). Classification of the types of androgenetic alopecia occurring in the female sex. Br J Dermatol, 97(3):247-254."),
    ("8", "Giovannucci E, et al. (1997). The CAG repeat within the androgen receptor gene and its relationship to prostate cancer. Proc Natl Acad Sci USA, 94(7):3320-3323."),
    ("9", "Yip L, et al. (2009). Gene-wide association study between the aromatase gene (CYP19A1) and female pattern hair loss. Br J Dermatol, 161(2):289-294."),
]

# Insert references after REFERENSI para
insert_pos = ref_para_index + 1
for num, ref_text in missing_refs:
    p = doc.paragraphs[insert_pos]._element
    new_p = doc.add_paragraph(f"{num}. {ref_text}")._element
    p.addprevious(new_p)
    insert_pos += 1

print(f"\n[OK] Added {len(missing_refs)} missing references:")
for num, ref_text in missing_refs:
    preview = ref_text[:70] if len(ref_text) > 70 else ref_text
    print(f"  [{num}] {preview}...")

# Save
doc.save('Folliscope_IEEE.docx')
print("\n[OK] File saved: Folliscope_IEEE.docx")
print("\n[MANUAL STEPS NEEDED]:")
print("  1. Re-number all references to maintain [1-9] order")
print("  2. Update citations in text [1,2,3...] to match new numbering")
print("  3. Remove/replace future-dated refs (2025, 2026)")
